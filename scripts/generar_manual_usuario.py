from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_cell_border(cell, **kwargs):
    """Helper para bordes de tabla"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(tag)
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcPr.append(element)
            element.set(docx.oxml.ns.qn('w:val'), 'single')
            element.set(docx.oxml.ns.qn('w:sz'), '4')
            element.set(docx.oxml.ns.qn('w:space'), '0')
            element.set(docx.oxml.ns.qn('w:color'), 'auto')

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(37, 99, 235)  # blue-600
    return heading

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(71, 85, 105)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    return p

def main():
    doc = Document()
    
    # Estilo por defecto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== PORTADA ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MANUAL DE USUARIO")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(37, 99, 235)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Sistema Empresarial TechSolutions S.A.")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Dylan Albrygth Caal Caal\n")
    run.font.size = Pt(14)
    run.bold = True
    run = info.add_run("Carnet: 202140273\n")
    run.font.size = Pt(12)
    run = info.add_run("Ingenieria en Ciencias y Sistemas\n")
    run.font.size = Pt(12)
    run = info.add_run("Curso de Practica Intermedia")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== CONTENIDO ==========
    add_heading_custom(doc, "1. Introduccion", level=1)
    doc.add_paragraph(
        "El Sistema Empresarial TechSolutions S.A. es una aplicacion web full-stack diseñada para la gestion integral "
        "de clientes, proyectos, tareas y usuarios dentro de una empresa. El sistema permite administrar operaciones "
        "corporativas de manera eficiente, con control de roles y acceso desde cualquier dispositivo con navegador web."
    )
    doc.add_paragraph(
        "Este manual tiene como objetivo guiar al usuario en el uso basico del sistema, explicando cada modulo "
        "y sus funcionalidades principales."
    )
    
    add_heading_custom(doc, "2. Requisitos", level=1)
    doc.add_paragraph("Para utilizar el sistema necesitas:", style='List Bullet')
    doc.add_paragraph("Conexion a internet.", style='List Bullet')
    doc.add_paragraph("Un navegador web moderno (Google Chrome, Mozilla Firefox, Microsoft Edge o Safari).", style='List Bullet')
    doc.add_paragraph("Para la app movil: un dispositivo Android con Android 8.0 o superior.", style='List Bullet')
    
    add_heading_custom(doc, "3. Acceso al Sistema", level=1)
    doc.add_paragraph("El sistema esta disponible en la siguiente URL:")
    add_code_block(doc, "https://proyecto-empresarial.vercel.app/")
    doc.add_paragraph(
        "Al ingresar, la primera pantalla que veras es la de inicio de sesion. Si aun no tienes una cuenta, "
        "puedes registrarte desde el enlace \"Registrate aqui\"."
    )
    
    add_heading_custom(doc, "4. Pantalla de Inicio de Sesion", level=1)
    doc.add_paragraph(
        "En esta pantalla debes ingresar tu correo electronico y contraseña. Tambien puedes alternar la visibilidad "
        "de la contraseña presionando el icono del ojo junto al campo."
    )
    doc.add_paragraph("Credenciales de administrador (demo):", style='List Bullet')
    add_code_block(doc, "Correo: admin@techsolutions.com\nContraseña: admin123")
    doc.add_paragraph("Credenciales de usuario normal (demo):", style='List Bullet')
    add_code_block(doc, "Correo: carlos.mendez@techsolutions.com\nContraseña: password123")
    
    add_heading_custom(doc, "5. Panel Principal (Dashboard)", level=1)
    doc.add_paragraph(
        "Despues de iniciar sesion, el sistema te lleva al Dashboard. Aqui veras un resumen visual de las estadisticas "
        "mas importantes del sistema."
    )
    doc.add_paragraph("Si eres Administrador, veras:", style='List Bullet')
    doc.add_paragraph("Total de clientes, proyectos, tareas pendientes, tareas completadas, usuarios registrados.", style='List Bullet 2')
    doc.add_paragraph("Lista de los ultimos proyectos creados.", style='List Bullet 2')
    doc.add_paragraph("Lista de las ultimas tareas registradas.", style='List Bullet 2')
    doc.add_paragraph("Si eres Usuario normal, veras:", style='List Bullet')
    doc.add_paragraph("Tus tareas asignadas, pendientes, en progreso y completadas.", style='List Bullet 2')
    doc.add_paragraph("Tus proyectos asignados.", style='List Bullet 2')
    
    add_heading_custom(doc, "6. Modulo Clientes", level=1)
    doc.add_paragraph(
        "Este modulo permite gestionar los clientes empresariales registrados en el sistema."
    )
    doc.add_paragraph("Ver clientes: Lista de todos los clientes con nombre, empresa, correo y estado.", style='List Bullet')
    doc.add_paragraph("Buscar clientes: Usa el buscador en la parte superior para filtrar por nombre, empresa o correo.", style='List Bullet')
    doc.add_paragraph("Ver detalle: Presiona \"Ver\" para ver la informacion completa de un cliente y sus proyectos asociados.", style='List Bullet')
    doc.add_paragraph("Crear cliente (solo Admin): Presiona \"+ Nuevo cliente\" y completa el formulario.", style='List Bullet')
    doc.add_paragraph("Editar cliente (solo Admin): Desde el detalle, presiona \"Editar\".", style='List Bullet')
    doc.add_paragraph("Eliminar cliente (solo Admin): Presiona \"Eliminar\" y confirma.", style='List Bullet')
    
    add_heading_custom(doc, "7. Modulo Proyectos", level=1)
    doc.add_paragraph(
        "Aqui se gestionan los proyectos asociados a los clientes. Cada proyecto muestra una barra de progreso visual "
        "calculada automaticamente a partir de las tareas completadas."
    )
    doc.add_paragraph("Ver proyectos: Tarjetas con nombre del proyecto, cliente, estado y porcentaje de avance.", style='List Bullet')
    doc.add_paragraph("Buscar proyectos: Filtra por nombre del proyecto, cliente o empresa.", style='List Bullet')
    doc.add_paragraph("Ver detalle: Muestra informacion completa, progreso, tareas asociadas y permite descargar un reporte PDF.", style='List Bullet')
    doc.add_paragraph("Descargar PDF: Desde el detalle del proyecto, presiona \"Descargar PDF\" para generar un reporte con toda la informacion.", style='List Bullet')
    doc.add_paragraph("Crear proyecto (solo Admin): Presiona \"+ Nuevo proyecto\".", style='List Bullet')
    doc.add_paragraph("Editar / Eliminar (solo Admin): Disponibles desde el detalle del proyecto.", style='List Bullet')
    
    add_heading_custom(doc, "8. Modulo Tareas", level=1)
    doc.add_paragraph(
        "Las tareas estan agrupadas por proyecto. Cada tarea tiene un titulo, prioridad, estado, responsable y fecha limite."
    )
    doc.add_paragraph("Ver tareas: Lista agrupada por proyecto con estadisticas de estado (pendientes, en progreso, completadas).", style='List Bullet')
    doc.add_paragraph("Buscar tareas: Filtra por titulo, proyecto o responsable.", style='List Bullet')
    doc.add_paragraph("Crear tarea: Presiona \"+ Nueva tarea\" junto a cada proyecto. Tambien puedes crear desde el boton general.", style='List Bullet')
    doc.add_paragraph("Ver historial: Presiona \"Historial\" para ver el timeline de fechas de la tarea.", style='List Bullet')
    doc.add_paragraph("Editar / Eliminar (solo Admin): Botones disponibles en cada tarea.", style='List Bullet')
    
    add_heading_custom(doc, "9. Modulo Usuarios (Solo Administrador)", level=1)
    doc.add_paragraph(
        "Este modulo solo esta disponible para usuarios con rol de Administrador. Permite gestionar los usuarios del sistema."
    )
    doc.add_paragraph("Ver usuarios: Lista con nombre, correo, rol y cantidad de tareas asignadas.", style='List Bullet')
    doc.add_paragraph("Buscar usuarios: Filtra por nombre o correo.", style='List Bullet')
    doc.add_paragraph("Ver tareas de usuario: Presiona \"Ver tareas\" para ver las tareas asignadas a un usuario especifico.", style='List Bullet')
    doc.add_paragraph("Eliminar usuario: Presiona \"Eliminar\" y confirma.", style='List Bullet')
    
    add_heading_custom(doc, "10. Perfil de Usuario", level=1)
    doc.add_paragraph(
        "Desde el menu lateral inferior, presiona \"Perfil\" para ver tu informacion personal. "
        "Aqui puedes ver tu nombre, correo y rol dentro del sistema."
    )
    
    add_heading_custom(doc, "11. Cerrar Sesion", level=1)
    doc.add_paragraph(
        "En el menu lateral inferior, presiona \"Cerrar sesion\" para salir del sistema. "
        "Se eliminara tu sesion y seras redirigido a la pantalla de login."
    )
    
    add_heading_custom(doc, "12. App Movil Android", level=1)
    doc.add_paragraph(
        "El sistema cuenta con una aplicacion movil para Android, desarrollada con Capacitor. "
        "La app carga la interfaz directamente desde Vercel, por lo que siempre tendras la version mas reciente."
    )
    add_heading_custom(doc, "12.1 Instalacion", level=2)
    doc.add_paragraph("Descarga el archivo app-debug.apk en tu celular Android.", style='List Bullet')
    doc.add_paragraph("Abre el archivo desde el gestor de archivos.", style='List Bullet')
    doc.add_paragraph("Si te pide \"Permitir fuentes desconocidas\", activa el permiso.", style='List Bullet')
    doc.add_paragraph("Presiona \"Instalar\" y espera a que termine.", style='List Bullet')
    doc.add_paragraph("La app aparecera en tu menu de aplicaciones con el nombre \"TechSolutions\".", style='List Bullet')
    add_heading_custom(doc, "12.2 Uso", level=2)
    doc.add_paragraph("Abre la app.", style='List Bullet')
    doc.add_paragraph("Ingresa tus credenciales (las mismas que en la version web).", style='List Bullet')
    doc.add_paragraph("Navega por los modulos igual que en la version web.", style='List Bullet')
    doc.add_paragraph("Requiere conexion a internet para funcionar.", style='List Bullet')
    
    add_heading_custom(doc, "13. Solucion de Problemas", level=1)
    doc.add_paragraph("No puedo iniciar sesion:", style='List Bullet')
    doc.add_paragraph("Verifica que tu correo y contraseña esten correctos. Usa el icono del ojo para ver la contraseña.", style='List Bullet 2')
    doc.add_paragraph("La pagina no carga:", style='List Bullet')
    doc.add_paragraph("Verifica tu conexion a internet. La URL debe ser: https://proyecto-empresarial.vercel.app/", style='List Bullet 2')
    doc.add_paragraph("No veo el boton de crear/editar/eliminar:", style='List Bullet')
    doc.add_paragraph("Esas funciones solo estan disponibles para el rol de Administrador.", style='List Bullet 2')
    doc.add_paragraph("La app movil no abre:", style='List Bullet')
    doc.add_paragraph("Asegurate de tener Android 8.0 o superior y conexion a internet.", style='List Bullet 2')
    
    add_heading_custom(doc, "14. Enlaces Importantes", level=1)
    p = doc.add_paragraph()
    p.add_run("Repositorio GitHub: ").bold = True
    p.add_run("https://github.com/dylanalbrygth2002-boop/Proyecto_Empresarial")
    p = doc.add_paragraph()
    p.add_run("Sistema en linea: ").bold = True
    p.add_run("https://proyecto-empresarial.vercel.app/")
    
    add_heading_custom(doc, "15. Credenciales de Prueba", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rol'
    hdr_cells[1].text = 'Correo'
    hdr_cells[2].text = 'Contraseña'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    users = [
        ("Administrador", "admin@techsolutions.com", "admin123"),
        ("Usuario", "carlos.mendez@techsolutions.com", "password123"),
        ("Usuario", "ana.garcia@techsolutions.com", "password123"),
        ("Usuario", "luis.rodriguez@techsolutions.com", "password123"),
        ("Usuario", "maria.lopez@techsolutions.com", "password123"),
        ("Usuario", "juan.perez@techsolutions.com", "password123"),
        ("Usuario", "sofia.torres@techsolutions.com", "password123"),
        ("Usuario", "diego.ramirez@techsolutions.com", "password123"),
        ("Usuario", "valentina.flores@techsolutions.com", "password123"),
        ("Usuario", "andres.morales@techsolutions.com", "password123"),
        ("Usuario", "camila.diaz@techsolutions.com", "password123"),
    ]
    
    for role, email, password in users:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = email
        row_cells[2].text = password
    
    doc.add_paragraph()
    doc.add_paragraph("---")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TechSolutions S.A. - Proyecto Integrador de Practica Intermedia\nDylan Albrygth Caal Caal - 202140273")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(148, 163, 184)
    
    doc.save("MANUAL_USUARIO.docx")
    print("Documento MANUAL_USUARIO.docx generado exitosamente.")

if __name__ == "__main__":
    main()
