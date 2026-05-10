from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(37, 99, 235)
    return heading

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(71, 85, 105)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    return p

def main():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== PORTADA ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ENTREGABLES PRACTICA INTERMEDIA")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(37, 99, 235)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Sistema Empresarial TechSolutions S.A.\nFull-Stack con Next.js, React y PostgreSQL")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Dylan Albrygth Caal Caal\n")
    run.font.size = Pt(14)
    run.bold = True
    run = info.add_run("Carnet: 202140273\n")
    run.font.size = Pt(12)
    run = info.add_run("Ingenieria en Ciencias y Sistemas\n")
    run.font.size = Pt(12)
    run = info.add_run("Curso de Practica Intermedia")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== INDICE ==========
    add_heading_custom(doc, "Indice de Entregables", level=1)
    items = [
        "1. Codigo Fuente Completo",
        "2. Repositorio GitHub",
        "3. URL del Sistema Desplegado",
        "4. Documentacion Tecnica (README.md)",
        "5. Manual Basico de Usuario",
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_page_break()
    
    # ========== 1. CODIGO FUENTE ==========
    add_heading_custom(doc, "1. Codigo Fuente Completo", level=1)
    doc.add_paragraph(
        "El codigo fuente completo del proyecto (frontend y backend) se encuentra en el repositorio de GitHub. "
        "El proyecto esta desarrollado con la arquitectura full-stack usando Next.js App Router, donde el backend "
        "se implementa mediante API Routes (Route Handlers) dentro del mismo proyecto."
    )
    
    add_heading_custom(doc, "1.1 Estructura del Codigo", level=2)
    doc.add_paragraph("Frontend (Next.js App Router):", style='List Bullet')
    doc.add_paragraph("app/ — Paginas y componentes React", style='List Bullet 2')
    doc.add_paragraph("components/ui/ — Componentes reutilizables (Button, Card, Input, Badge, Alert)", style='List Bullet 2')
    doc.add_paragraph("components/layout/ — AppShell, AppSidebar, AuthProvider", style='List Bullet 2')
    doc.add_paragraph("Backend (API Routes):", style='List Bullet')
    doc.add_paragraph("app/api/auth/ — Autenticacion (login, register, logout, me)", style='List Bullet 2')
    doc.add_paragraph("app/api/clientes/ — CRUD de clientes", style='List Bullet 2')
    doc.add_paragraph("app/api/proyectos/ — CRUD de proyectos + reportes PDF", style='List Bullet 2')
    doc.add_paragraph("app/api/tareas/ — CRUD de tareas", style='List Bullet 2')
    doc.add_paragraph("app/api/usuarios/ — Gestion de usuarios", style='List Bullet 2')
    doc.add_paragraph("app/api/dashboard/ — Resumen estadistico", style='List Bullet 2')
    doc.add_paragraph("Base de datos:", style='List Bullet')
    doc.add_paragraph("prisma/schema.prisma — Modelo de datos (User, Client, Project, Task)", style='List Bullet 2')
    doc.add_paragraph("prisma/seed.ts — Datos demo", style='List Bullet 2')
    
    add_heading_custom(doc, "1.2 Tecnologias Utilizadas", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "Categoria"
    hdr[1].text = "Tecnologia"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    
    data = [
        ("Framework", "Next.js 16 (App Router)"),
        ("Frontend", "React 19, TypeScript, TailwindCSS 4"),
        ("Backend", "Next.js API Routes (Route Handlers)"),
        ("ORM", "Prisma 6"),
        ("Base de datos", "PostgreSQL (local / Railway)"),
        ("Runtime", "Bun"),
        ("Autenticacion", "JWT + localStorage"),
        ("Validacion", "Zod"),
        ("Reportes", "jsPDF + jspdf-autotable"),
        ("App movil", "Capacitor 8 + Android"),
        ("Despliegue", "Vercel (frontend) + Railway (PostgreSQL)"),
    ]
    for cat, tech in data:
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = tech
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Ubicacion del codigo fuente: ").bold = True
    p.add_run("https://github.com/dylanalbrygth2002-boop/Proyecto_Empresarial")
    
    doc.add_page_break()
    
    # ========== 2. REPOSITORIO GITHUB ==========
    add_heading_custom(doc, "2. Repositorio GitHub", level=1)
    doc.add_paragraph(
        "El proyecto se encuentra alojado en GitHub con control de versiones mediante Git. "
        "El repositorio incluye todo el historial de commits, documentacion, scripts automatizados y configuraciones."
    )
    p = doc.add_paragraph()
    p.add_run("URL del repositorio: ").bold = True
    p.add_run("https://github.com/dylanalbrygth2002-boop/Proyecto_Empresarial")
    
    add_heading_custom(doc, "2.1 Historial de Commits Relevantes", level=2)
    commits = [
        ("308b718", "docs: actualizar README.md con documentacion completa"),
        ("ef2e37e", "Fix: texto que se salia de cuadros en vistas moviles"),
        ("25186b7", "Animaciones, toggle contraseña y sidebar movil sin scroll"),
        ("687ab99", "Diseño: color aplicado a TODAS las cards"),
        ("f92491e", "Diseño: cards con color, headers con gradiente"),
        ("c8a6b9f", "Fix: boton hamburguesa ya no tapa logo en movil"),
        ("a46210e", "Fix: traducir Dashboard a Inicio"),
        ("5fd7164", "Rediseño visual completo"),
    ]
    for hash, msg in commits:
        doc.add_paragraph(f"{hash} — {msg}", style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 3. URL DESPLEGADA ==========
    add_heading_custom(doc, "3. URL del Sistema Desplegado", level=1)
    doc.add_paragraph(
        "El sistema esta desplegado en Vercel conectado al repositorio de GitHub. "
        "Cada push a la rama main despliega automaticamente la ultima version. "
        "La base de datos PostgreSQL esta alojada en Railway."
    )
    p = doc.add_paragraph()
    p.add_run("URL de acceso: ").bold = True
    p.add_run("https://proyecto-empresarial.vercel.app/")
    p = doc.add_paragraph()
    p.add_run("Plataforma de despliegue: ").bold = True
    p.add_run("Vercel")
    p = doc.add_paragraph()
    p.add_run("Base de datos: ").bold = True
    p.add_run("PostgreSQL en Railway")
    p = doc.add_paragraph()
    p.add_run("Configuracion de build: ").bold = True
    p.add_run("bunx prisma migrate deploy && bun run build (vercel.json)")
    
    add_heading_custom(doc, "3.1 Variables de Entorno en Produccion", level=2)
    add_code_block(doc, "DATABASE_URL=postgresql://... (Railway)\nJWT_SECRET=tu-secreto-jwt\nNODE_ENV=production")
    
    doc.add_page_break()
    
    # ========== 4. DOCUMENTACION TECNICA ==========
    add_heading_custom(doc, "4. Documentacion Tecnica", level=1)
    doc.add_paragraph(
        "La documentacion tecnica completa se encuentra en el archivo README.md del repositorio. "
        "A continuacion se presenta un resumen de los aspectos tecnicos mas importantes."
    )
    
    add_heading_custom(doc, "4.1 Modelo de Datos (Prisma Schema)", level=2)
    doc.add_paragraph("Entidades principales:", style='List Bullet')
    doc.add_paragraph("User — id, name, email, password, role (ADMIN/USER), createdAt", style='List Bullet 2')
    doc.add_paragraph("Client — id, name, email, phone, company, status, createdAt", style='List Bullet 2')
    doc.add_paragraph("Project — id, name, description, status, startDate, endDate, clientId", style='List Bullet 2')
    doc.add_paragraph("Task — id, title, description, status, priority, dueDate, projectId, responsibleId", style='List Bullet 2')
    
    add_heading_custom(doc, "4.2 Endpoints API", level=2)
    endpoints = [
        ("POST /api/auth/register", "Registro de usuarios"),
        ("POST /api/auth/login", "Inicio de sesion con JWT"),
        ("POST /api/auth/logout", "Cerrar sesion"),
        ("GET /api/auth/me", "Obtener usuario actual"),
        ("GET /api/clientes", "Listar clientes"),
        ("POST /api/clientes", "Crear cliente"),
        ("GET /api/proyectos", "Listar proyectos con progreso"),
        ("POST /api/proyectos", "Crear proyecto"),
        ("GET /api/proyectos/[id]/reporte", "Generar reporte PDF"),
        ("GET /api/tareas", "Listar tareas filtradas por rol"),
        ("POST /api/tareas", "Crear tarea"),
        ("GET /api/usuarios", "Listar usuarios (admin)"),
        ("GET /api/dashboard/resumen", "Estadisticas del dashboard"),
    ]
    for endpoint, desc in endpoints:
        doc.add_paragraph(f"{endpoint} — {desc}", style='List Bullet')
    
    add_heading_custom(doc, "4.3 Seguridad", level=2)
    doc.add_paragraph("Autenticacion: JWT almacenado en localStorage", style='List Bullet')
    doc.add_paragraph("Autorizacion: Headers X-User-Id y X-User-Role en cada request", style='List Bullet')
    doc.add_paragraph("Proteccion de APIs: requireAdmin middleware en POST/PUT/DELETE", style='List Bullet')
    doc.add_paragraph("Validacion: Zod en todos los endpoints", style='List Bullet')
    doc.add_paragraph("Contraseñas: bcryptjs con hash seguro", style='List Bullet')
    
    add_heading_custom(doc, "4.4 App Movil", level=2)
    doc.add_paragraph("Framework: Capacitor 8", style='List Bullet')
    doc.add_paragraph("Estrategia: Remote URL (carga desde Vercel)", style='List Bullet')
    doc.add_paragraph("Plugins: status-bar, splash-screen, app, keyboard", style='List Bullet')
    doc.add_paragraph("Iconos y splash: Generados automaticamente desde logo base", style='List Bullet')
    doc.add_paragraph("APK de debug: android/app/build/outputs/apk/debug/app-debug.apk", style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 5. MANUAL DE USUARIO ==========
    add_heading_custom(doc, "5. Manual Basico de Usuario", level=1)
    doc.add_paragraph(
        "Este apartado describe el uso basico del sistema para los usuarios finales."
    )
    
    add_heading_custom(doc, "5.1 Acceso al Sistema", level=2)
    doc.add_paragraph("Abre tu navegador y ve a: https://proyecto-empresarial.vercel.app/")
    doc.add_paragraph("Ingresa tu correo y contraseña. Presiona el icono del ojo para ver/ocultar la contraseña.")
    
    add_heading_custom(doc, "5.2 Credenciales de Prueba", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "Rol"
    hdr[1].text = "Correo"
    hdr[2].text = "Contraseña"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    
    test_users = [
        ("Administrador", "admin@techsolutions.com", "admin123"),
        ("Usuario", "carlos.mendez@techsolutions.com", "password123"),
        ("Usuario", "ana.garcia@techsolutions.com", "password123"),
    ]
    for role, email, pwd in test_users:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = email
        row[2].text = pwd
    
    add_heading_custom(doc, "5.3 Navegacion por Modulos", level=2)
    doc.add_paragraph("Dashboard: Resumen visual con estadisticas y listados recientes.", style='List Bullet')
    doc.add_paragraph("Clientes: Lista de clientes empresariales. Admin puede crear, editar y eliminar.", style='List Bullet')
    doc.add_paragraph("Proyectos: Tarjetas con avance visual. Admin puede gestionar todo. Incluye reporte PDF.", style='List Bullet')
    doc.add_paragraph("Tareas: Agrupadas por proyecto. Incluye historial de fechas.", style='List Bullet')
    doc.add_paragraph("Usuarios: Solo visible para admin. Permite ver tareas asignadas de cada usuario.", style='List Bullet')
    doc.add_paragraph("Perfil: Informacion del usuario autenticado.", style='List Bullet')
    
    add_heading_custom(doc, "5.4 App Movil Android", level=2)
    doc.add_paragraph("Descarga e instala el archivo app-debug.apk en tu celular Android.", style='List Bullet')
    doc.add_paragraph("Requiere Android 8.0+ y conexion a internet.", style='List Bullet')
    doc.add_paragraph("Usa las mismas credenciales que la version web.", style='List Bullet')
    doc.add_paragraph("Al cerrar y abrir la app, se carga la ultima version automaticamente.", style='List Bullet')
    
    add_heading_custom(doc, "5.5 Solucion de Problemas", level=2)
    doc.add_paragraph("No puedo iniciar sesion: Verifica correo y contraseña. Usa el icono del ojo.", style='List Bullet')
    doc.add_paragraph("Pagina no carga: Verifica conexion a internet.", style='List Bullet')
    doc.add_paragraph("No veo botones de crear/editar/eliminar: Esas funciones son solo para Administrador.", style='List Bullet')
    
    # ========== FOOTER ==========
    doc.add_paragraph()
    doc.add_paragraph("---")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TechSolutions S.A. - Proyecto Integrador de Practica Intermedia\nDylan Albrygth Caal Caal - 202140273\nIngenieria en Ciencias y Sistemas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(148, 163, 184)
    
    doc.save("ENTREGABLES_PRACTICA.docx")
    print("Documento ENTREGABLES_PRACTICA.docx generado exitosamente.")

if __name__ == "__main__":
    main()
